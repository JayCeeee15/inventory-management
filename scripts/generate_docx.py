from __future__ import annotations

from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError as exc:  # pragma: no cover
    raise SystemExit(
        "python-docx is required to run this script. Install it with: pip install python-docx"
    ) from exc


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
OUTPUT_PATH = DOCS_DIR / "Hospital_Inventory_System_Documentation.docx"


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_heading_paragraph(document: Document, heading: str, body: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    heading_run = paragraph.add_run(f"{heading}: ")
    heading_run.bold = True
    paragraph.add_run(body)


def add_title_page(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Hospital Inventory Management System")
    run.bold = True
    run.font.size = Pt(24)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Tech Stack: Angular, Express.js, MySQL")
    subtitle_run.font.size = Pt(14)

    document.add_paragraph("")

    prepared_by = document.add_paragraph()
    prepared_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prepared_by.add_run("Prepared by:\n\n______________________________")

    doc_date = document.add_paragraph()
    doc_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_date.add_run(f"Date: {date.today().strftime('%B %d, %Y')}")

    document.add_page_break()


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    add_title_page(document)

    document.add_heading("1) System Overview", level=1)
    add_bullets(
        document,
        [
            "The Hospital Inventory Management System tracks medicine, supply, and equipment stocks for a hospital environment.",
            "It supports stock receiving, walk-in sales, patient issue transactions, transaction history, dashboard monitoring, and profile management.",
            "Admin users manage stock intake, master data, reports, and overall control of the system.",
            "Staff users focus on daily operational transactions such as walk-in sales and patient issue postings.",
            "High-level workflow: Admin receives stock -> inventory quantities increase -> Staff sells or issues stock -> dashboards and transaction history update."
        ],
    )

    document.add_heading("2) Key Modules", level=1)
    add_heading_paragraph(document, "Dashboard", "Summary cards, stock graph, latest activities, and recent transactions for quick monitoring.")
    add_heading_paragraph(document, "Item Master", "Create, review, update, and organize medicine, supply, and equipment records.")
    add_heading_paragraph(document, "Categories", "Manage item groupings such as pharmacy stock, PPE, emergency medicines, and ward consumables.")
    add_heading_paragraph(document, "Receiving / Stock In", "Admin posts stock intake to increase stock-on-hand and record a RECEIVE stock ledger entry.")
    add_heading_paragraph(document, "Walk-in Sale", "Staff records over-the-counter patient purchases and deducts stock in real time.")
    add_heading_paragraph(document, "Patient Issue", "Staff issues stocks to departments or patients while validating available quantity.")
    add_heading_paragraph(document, "Transaction History", "Provides a searchable audit trail of stock movements and business transactions.")
    add_heading_paragraph(document, "Profile & Avatar Upload", "Lets users update their profile information and profile photo.")
    add_heading_paragraph(document, "Role-Based Access", "Separates Admin controls from Staff actions for safer and clearer workflows.")

    document.add_heading("3) Why Angular is Important for This System", level=1)
    add_heading_paragraph(document, "Component-based UI", "Angular lets this project split the interface into reusable parts such as dashboard cards, tables, forms, modals, and side navigation.")
    add_heading_paragraph(document, "Routing", "Angular routing makes it easy to move between home, login, admin dashboard, and employee dashboard while protecting each route with guards.")
    add_heading_paragraph(document, "Forms", "Reactive Forms give structured validation for walk-in sale, patient issue, add item, receiving, and profile update forms.")
    add_heading_paragraph(document, "Services + HttpClient", "Angular services organize API calls to the Express backend so components stay focused on the UI.")
    add_heading_paragraph(document, "State Management", "RxJS, Signals, or NgRx can coordinate dashboard refreshes, shared loading states, and transaction updates across pages.")
    add_heading_paragraph(document, "Performance", "Angular offers structured change detection, async patterns, and reusable rendering strategies that matter in data-heavy dashboards.")
    add_heading_paragraph(document, "Maintainability", "A clear Angular structure helps the system grow without turning into a single large, difficult-to-maintain file.")

    document.add_heading("4) What I Should Learn in Angular (Learning Roadmap)", level=1)
    add_numbered(
        document,
        [
            "Angular fundamentals: components, templates, interpolation, property binding, event binding, and structural directives.",
            "Routing and guards: understand route navigation, auth guards, and role guards so Admin and Staff views stay separated.",
            "Reactive Forms: learn FormGroup, FormArray, validation, and form submission for walk-in sale, patient issue, add item, and receiving forms.",
            "HttpClient and interceptors: learn how Angular talks to Express, attaches JWT tokens, handles API errors, and centralizes retry behavior.",
            "Observables and RxJS basics: understand subscribe, switchMap, forkJoin, finalize, takeUntil, and retry because many dashboard and form flows depend on them.",
            "UI framework best practices: learn Angular Material or PrimeNG patterns for cards, tables, buttons, dialogs, and forms.",
            "Change detection and loading issues: study why data may appear only after a click, and how lifecycle hooks, finalize, async pipe, and ChangeDetectorRef help fix it.",
            "Tables and filtering: practice pagination, searching, filtering, and empty states for Item Master and Transaction History.",
            "File upload: learn multipart/form-data upload for the profile photo feature.",
            "Deployment basics: understand build configurations, environment files, and how the Angular frontend connects to Express and MySQL in different environments."
        ],
    )

    document.add_heading("5) Angular vs Other Programming Languages/Frameworks", level=1)
    document.add_paragraph(
        "Angular is a frontend framework built with TypeScript. It is not a general programming language by itself. "
        "It focuses on building browser-based user interfaces, while backend languages and frameworks handle server logic, databases, authentication, and business rules."
    )

    comparison_headers = ["Comparison", "Primary Focus", "Strengths", "Tradeoffs", "What It Means for This Project"]
    comparison_rows = [
        [
            "Angular vs plain JavaScript/jQuery",
            "Structured frontend framework vs manual DOM scripting",
            "Angular gives reusable components, routing, forms, services, and cleaner large-project structure",
            "Angular has a bigger learning curve than simple scripts",
            "Better for a hospital system because the UI has many forms, roles, dashboards, and shared components",
        ],
        [
            "Angular vs React",
            "Full framework vs UI library",
            "Angular includes routing, forms, dependency injection, and conventions out of the box",
            "React is often more flexible, but usually needs more package decisions",
            "Angular is useful when you want an opinionated enterprise-style structure for a business system",
        ],
        [
            "Angular vs Vue",
            "Both are frontend frameworks for building UIs",
            "Angular is stronger in large enterprise structure and built-in tooling",
            "Vue can feel lighter and easier at the start",
            "Angular fits well when the project needs formal structure, guards, modularity, and team scaling",
        ],
        [
            "Angular vs Node/Express, PHP, Java, C#",
            "Frontend vs backend responsibilities",
            "Angular handles screens, forms, navigation, and user interaction",
            "It does not replace the backend or the database",
            "Express and MySQL handle API logic, transactions, stock rules, and persistent storage while Angular displays and sends data",
        ],
    ]

    comparison_table = document.add_table(rows=1, cols=len(comparison_headers))
    comparison_table.style = "Table Grid"
    for index, header in enumerate(comparison_headers):
        comparison_table.rows[0].cells[index].text = header
    for row in comparison_rows:
        cells = comparison_table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value

    document.add_heading("6) Common Issues in My System and How Angular Helps Fix Them", level=1)
    add_heading_paragraph(document, "Loading stuck on 'Saving...'", "Use finalize() to reset loading flags and ensure the UI exits the saving state even when an error happens.")
    add_heading_paragraph(document, "Retry button not working", "Re-trigger the observable or API method and reset local loading and error state before retrying.")
    add_heading_paragraph(document, "Sidenav stretching when collapsing", "Keep icon containers fixed and only animate label opacity or position through CSS.")
    add_heading_paragraph(document, "Real-time dashboard updates", "Use a shared service, Subject, Signal, or direct re-fetch after a successful transaction post.")
    add_heading_paragraph(document, "Stock mismatch after sale", "Always refresh stock values from the API after the backend transaction commits instead of trusting stale UI values.")

    document.add_heading("7) Best Practices Checklist", level=1)
    add_bullets(
        document,
        [
            "Handle API responses carefully and return clear HTTP status codes.",
            "Enforce role-based access both in the Angular UI and in the Express backend.",
            "Use database transactions for receiving, sale, and patient issue posting.",
            "Keep audit logs and stock ledger entries for every stock-changing action.",
            "Prevent negative stock through defensive validation before commit.",
            "Use a clean Angular folder structure with core, shared, and feature modules.",
        ],
    )

    document.add_heading("8) Appendix", level=1)
    document.add_heading("Suggested Angular Folder Structure", level=2)
    folder_headers = ["Folder", "Purpose"]
    folder_rows = [
        ["src/app/core", "Auth service, guards, interceptors, app-wide services, and API integration."],
        ["src/app/shared", "Reusable models, utilities, and presentational components."],
        ["src/app/features/home", "Landing page and public-facing content."],
        ["src/app/features/auth", "Login and signup screens."],
        ["src/app/features/dashboard", "Admin dashboard and admin-only modules."],
        ["src/app/features/employee-dashboard", "Staff dashboard and staff transaction flows."],
        ["src/app/features/products", "Item Master list and product form components."],
        ["src/app/features/categories", "Category list and category form components."],
        ["src/app/features/transactions", "Walk-in sale, patient issue, receiving, and transaction history components."],
    ]
    folder_table = document.add_table(rows=1, cols=2)
    folder_table.style = "Table Grid"
    for index, header in enumerate(folder_headers):
        folder_table.rows[0].cells[index].text = header
    for row in folder_rows:
        cells = folder_table.add_row().cells
        cells[0].text = row[0]
        cells[1].text = row[1]

    document.add_heading("Example REST Endpoints", level=2)
    add_bullets(
        document,
        [
            "POST /api/auth/login - authenticate a user",
            "GET /api/inventory/dashboard/summary - load dashboard metrics",
            "GET /api/inventory/products - load item master data",
            "POST /api/inventory/stock/receive - admin stock receiving",
            "POST /api/inventory/sales - staff walk-in sale posting",
            "POST /api/inventory/patient-issues - staff patient issue posting",
            "GET /api/inventory/stock/movements - transaction history and stock ledger view",
            "DELETE /api/inventory/transactions/:movementId - admin transaction delete",
        ],
    )

    return document


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
